from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "finance-report.md"
OUTPUT = ROOT / "docs" / "财务Excel自动化处理项目汇报.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
TEXT = "1F2933"
MUTED = "5B6770"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_east_asia_font(run, font_name: str = "Hiragino Sans GB") -> None:
    run.font.name = "Hiragino Sans GB"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Hiragino Sans GB")
    r_fonts.set(qn("w:hAnsi"), "Hiragino Sans GB")


def style_run(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    set_east_asia_font(run)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Hiragino Sans GB"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("财务 Excel 自动化处理项目汇报")
    style_run(run, size=9, color=MUTED)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("财务 Excel 自动化处理项目汇报")
    style_run(run, size=24, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("面向财务团队与管理层的流程、架构与效率提升说明")
    style_run(run, size=12, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(callout, [7200])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    set_cell_margins(cell, top=160, bottom=160, start=220, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("核心结论：将月度 Excel 报表中重复的筛选、汇总、计算和填表流程自动化，预计单次处理效率提升约 70%-90%，常规重复性操作错误降低约 60%-80%。")
    style_run(r, size=12, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    run = meta.add_run("文档用途：内部汇报 / 流程说明 / 自动化项目价值评估")
    style_run(run, size=10, color=MUTED)

    doc.add_section(WD_SECTION.NEW_PAGE)


def split_inline_code(text: str):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            yield part[1:-1], True
        else:
            yield part, False


def add_text_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    for value, is_code in split_inline_code(text):
        run = p.add_run(value)
        style_run(run, size=10 if is_code else None, color=TEXT)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    if style in ("Heading 1", "Heading 2", "Heading 3"):
        p.paragraph_format.keep_with_next = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    for value, is_code in split_inline_code(text):
        run = p.add_run(value)
        style_run(run, size=10 if is_code else None, color=TEXT)


def add_numbered(doc: Document, number: str, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    for value, is_code in split_inline_code(text):
        run = p.add_run(value)
        style_run(run, size=10 if is_code else None, color=TEXT)


def parse_table(lines: list[str]):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(c.replace(":", "").replace("-", "")) == set() for c in rows[1]):
        rows.pop(1)
    return rows


def choose_widths(headers: list[str], cols: int) -> list[int]:
    if cols == 2:
        return [2600, 6760]
    if cols == 3:
        return [2100, 3600, 3660]
    if cols == 4:
        if headers and headers[0] == "工作环节":
            return [2200, 1900, 1900, 3360]
        return [1800, 2520, 2520, 2520]
    return [9360 // cols] * cols


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = choose_widths(rows[0], len(rows[0]))
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx > 0 and re.search(r"\\d|%|分钟|小时", value):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value.replace("<br>", "\n"))
            style_run(run, size=10, bold=(r_idx == 0), color=DARK_BLUE if r_idx == 0 else TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [9000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("334155")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_footer(doc)
    add_title_page(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith("# "):
            add_text_paragraph(doc, stripped[2:], "Heading 1")
        elif stripped.startswith("## "):
            add_text_paragraph(doc, stripped[3:], "Heading 1")
        elif stripped.startswith("### "):
            add_text_paragraph(doc, stripped[4:], "Heading 2")
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
        elif re.match(r"^\\d+\\.\\s+", stripped):
            add_numbered(doc, "", re.sub(r"^\\d+\\.\\s+", "", stripped))
        else:
            add_text_paragraph(doc, stripped)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
